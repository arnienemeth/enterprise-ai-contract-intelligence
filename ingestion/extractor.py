"""Text extraction ("OCR") layer.

Turns the raw bytes of an uploaded file into plain text, regardless of format.
Follows the same dispatcher pattern as the vector store (Azure ↔ local):

  * If AZURE_DOCINTEL_ENDPOINT is set → Azure AI Document Intelligence.
    Handles real-world inputs the local libraries can't: scanned PDFs, photos
    of contracts, images, and complex layouts (tables, multi-column). This is
    the production path for "real" contracts.

  * Otherwise → local libraries (no extra cloud resource needed):
        - PDF  : pdfplumber (falls back to pypdf)
        - DOCX : python-docx
        - TXT  : utf-8 / latin-1 decode
    Good enough for *digital* (text-based) PDF/DOCX. A scanned/image-only PDF
    has no embedded text, so the local path returns little/nothing — that is the
    signal to configure Document Intelligence.

Public API:
    extract_text(file_bytes: bytes, filename: str) -> str
    ExtractionError               (raised on unrecoverable failures)
    SUPPORTED_EXTENSIONS          (set of handled extensions)

The function never calls the network unless the Document Intelligence path is
active, so it's cheap and safe to unit-test.
"""

import os

# Digital formats handled locally; Document Intelligence extends this to scans.
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

# Heuristic: if the local extractor yields fewer characters than this for a PDF,
# the file is almost certainly a scan/image (no embedded text layer).
_MIN_CHARS_FOR_TEXT_PDF = 20


class ExtractionError(Exception):
    """Raised when a file cannot be turned into usable text."""


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


# ---------------------------------------------------------------------------
# Local decoders
# ---------------------------------------------------------------------------

def _extract_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    # Last resort: never blow up on a stray byte.
    return data.decode("utf-8", errors="replace")


def _extract_pdf_local(data: bytes) -> str:
    """Digital-PDF text extraction. pdfplumber first (better layout), pypdf as fallback."""
    import io

    text = ""

    # 1) pdfplumber — preserves reading order / spacing better.
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [(page.extract_text() or "") for page in pdf.pages]
        text = "\n\n".join(pages).strip()
    except Exception:
        text = ""

    # 2) pypdf fallback if pdfplumber found nothing.
    if len(text) < _MIN_CHARS_FOR_TEXT_PDF:
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            pages = [(p.extract_text() or "") for p in reader.pages]
            alt = "\n\n".join(pages).strip()
            if len(alt) > len(text):
                text = alt
        except Exception:
            pass

    return text


def _extract_docx_local(data: bytes) -> str:
    import io

    from docx import Document

    document = Document(io.BytesIO(data))

    parts = [p.text for p in document.paragraphs]

    # Include table cell text — contracts often put terms in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


# ---------------------------------------------------------------------------
# Azure AI Document Intelligence (production OCR)
# ---------------------------------------------------------------------------

def _extract_with_document_intelligence(data: bytes) -> str:
    """Extract text via Azure AI Document Intelligence prebuilt-read (OCR).

    Handles scanned PDFs, images and complex layouts. Imported lazily so the app
    runs without the SDK installed when the feature is off.

    Config:
        AZURE_DOCINTEL_ENDPOINT   https://<resource>.cognitiveservices.azure.com/
        AZURE_DOCINTEL_KEY        resource key
        AZURE_DOCINTEL_MODEL      model id (default: "prebuilt-read")
    """
    endpoint = os.getenv("AZURE_DOCINTEL_ENDPOINT")
    key = os.getenv("AZURE_DOCINTEL_KEY")
    model = os.getenv("AZURE_DOCINTEL_MODEL", "prebuilt-read")

    if not endpoint or not key:
        raise ExtractionError(
            "AZURE_DOCINTEL_ENDPOINT is set but AZURE_DOCINTEL_KEY is missing."
        )

    from azure.core.credentials import AzureKeyCredential
    from azure.ai.documentintelligence import DocumentIntelligenceClient

    client = DocumentIntelligenceClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )

    poller = client.begin_analyze_document(model, body=data)
    result = poller.result()

    # `content` is the full concatenated text of the document.
    return (getattr(result, "content", "") or "").strip()


def _document_intelligence_enabled() -> bool:
    return bool(os.getenv("AZURE_DOCINTEL_ENDPOINT"))


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, filename: str = "document") -> str:
    """Extract plain text from an uploaded file's bytes.

    Strategy:
      * .txt / .md            → decoded locally (Document Intelligence never needed)
      * .pdf                  → local digital extraction; if that yields no text
                                and Document Intelligence is configured, fall
                                back to OCR (scans)
      * .docx                 → local; Document Intelligence as fallback if empty
      * anything else         → Document Intelligence if enabled, else error

    Raises ExtractionError if no text can be produced.
    """
    if file_bytes is None or len(file_bytes) == 0:
        raise ExtractionError(f"Empty file: {filename!r}")

    ext = _ext(filename)

    # Plain text: no OCR ever required.
    if ext in (".txt", ".md"):
        text = _extract_txt(file_bytes)
        if not text.strip():
            raise ExtractionError(f"No text content in {filename!r}.")
        return text

    # PDF: try local digital extraction first (cheap, no network).
    if ext == ".pdf":
        text = _extract_pdf_local(file_bytes)
        if len(text) >= _MIN_CHARS_FOR_TEXT_PDF:
            return text
        # Likely a scan → use Document Intelligence if available.
        if _document_intelligence_enabled():
            text = _extract_with_document_intelligence(file_bytes)
            if text.strip():
                return text
        raise ExtractionError(
            f"{filename!r} appears to be a scanned/image PDF with no extractable "
            "text. Configure Azure AI Document Intelligence "
            "(AZURE_DOCINTEL_ENDPOINT / AZURE_DOCINTEL_KEY) to OCR it."
        )

    # DOCX: local first, Document Intelligence as fallback.
    if ext == ".docx":
        try:
            text = _extract_docx_local(file_bytes)
        except Exception as e:
            text = ""
            if not _document_intelligence_enabled():
                raise ExtractionError(f"Could not read DOCX {filename!r}: {e}")
        if text.strip():
            return text
        if _document_intelligence_enabled():
            text = _extract_with_document_intelligence(file_bytes)
            if text.strip():
                return text
        raise ExtractionError(f"No text content in {filename!r}.")

    # Unknown / image formats: only Document Intelligence can handle these.
    if _document_intelligence_enabled():
        text = _extract_with_document_intelligence(file_bytes)
        if text.strip():
            return text
        raise ExtractionError(f"No text extracted from {filename!r}.")

    raise ExtractionError(
        f"Unsupported file type {ext!r} for {filename!r}. Supported locally: "
        f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}. Enable Azure AI Document "
        "Intelligence to handle scans and images."
    )
